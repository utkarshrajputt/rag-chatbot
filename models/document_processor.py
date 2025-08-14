"""
Document Processing Module
Handles PDF and DOCX file ingestion with improved text extraction and cleaning.
"""

import PyPDF2
from docx import Document
import re
import os
import logging
from typing import Optional, List

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document ingestion and text extraction for PDF and DOCX files."""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx']
    
    def clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text (str): Raw extracted text
            
        Returns:
            str: Cleaned text
        """
        # Remove non-printable characters except newlines
        text = re.sub(r'[^\x20-\x7E\n]', ' ', text)
        
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive newlines
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text.strip()
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """
        Extract text from PDF file.
        
        Args:
            pdf_path (str): Path to PDF file
            
        Returns:
            str: Extracted text
            
        Raises:
            ValueError: If PDF extraction fails
        """
        try:
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        extracted = page.extract_text()
                        if extracted:
                            text += extracted + "\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
                
                if not text.strip():
                    raise ValueError("No text could be extracted from the PDF")
                
                return self.clean_text(text)
                
        except Exception as e:
            logger.error(f"PDF extraction failed for {pdf_path}: {e}")
            raise ValueError(f"Failed to read PDF: {str(e)}")
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            docx_path (str): Path to DOCX file
            
        Returns:
            str: Extracted text
            
        Raises:
            ValueError: If DOCX extraction fails
        """
        try:
            doc = Document(docx_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            if not text.strip():
                raise ValueError("No text could be extracted from the DOCX file")
            
            return self.clean_text(text)
            
        except Exception as e:
            logger.error(f"DOCX extraction failed for {docx_path}: {e}")
            raise ValueError(f"Failed to read DOCX: {str(e)}")
    
    def process_document(self, file_path: str) -> str:
        """
        Process a document and extract its text content.
        
        Args:
            file_path (str): Path to the document file
            
        Returns:
            str: Extracted and cleaned text
            
        Raises:
            ValueError: If file format is unsupported or processing fails
        """
        if not os.path.exists(file_path):
            raise ValueError(f"File not found: {file_path}")
        
        file_ext = os.path.splitext(file_path.lower())[1]
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported formats: {self.supported_formats}")
        
        logger.info(f"Processing document: {file_path}")
        
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            return self.extract_text_from_docx(file_path)
    
    def validate_file(self, file_path: str, max_size_mb: int = 16) -> bool:
        """
        Validate file size and format.
        
        Args:
            file_path (str): Path to file
            max_size_mb (int): Maximum file size in MB
            
        Returns:
            bool: True if file is valid
        """
        if not os.path.exists(file_path):
            return False
        
        # Check file size
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        if file_size_mb > max_size_mb:
            logger.warning(f"File size ({file_size_mb:.2f} MB) exceeds limit ({max_size_mb} MB)")
            return False
        
        # Check file format
        file_ext = os.path.splitext(file_path.lower())[1]
        if file_ext not in self.supported_formats:
            logger.warning(f"Unsupported file format: {file_ext}")
            return False
        
        return True
